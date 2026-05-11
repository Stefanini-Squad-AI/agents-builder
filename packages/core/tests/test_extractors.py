"""Tests for every extractor + the registry + the truncation rule.

PDF and DOCX fixtures are generated on the fly via reportlab and
python-docx so the repo doesn't ship binary blobs.
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pytest
from app.extractors import (
    EXTRACTORS,
    TRUNCATION_THRESHOLD_BYTES,
    NoExtractorError,
    select_extractor,
    truncate_markdown,
)
from app.extractors.code_extractor import CodeExtractor
from app.extractors.csv_extractor import MAX_ROWS, CsvExtractor
from app.extractors.docx_extractor import DocxExtractor
from app.extractors.markdown_extractor import MarkdownExtractor
from app.extractors.pdf_extractor import PdfExtractor

# ---------------------------------------------------------------------------
# Truncation
# ---------------------------------------------------------------------------


def test_truncate_passes_short_content_through() -> None:
    content = "hello world"
    out, was = truncate_markdown(content)
    assert out == content
    assert was is False


def test_truncate_applies_at_threshold() -> None:
    # 2 MB of ASCII — bytes == chars in UTF-8.
    big = "x" * (2 * TRUNCATION_THRESHOLD_BYTES)
    out, was = truncate_markdown(big)
    assert was is True
    assert "truncated" in out
    assert "bytes omitted" in out
    # Head and tail were retained; total size of head+tail kept near
    # the configured budget.
    assert out.startswith("x" * 100)
    assert out.endswith("x" * 100)


# ---------------------------------------------------------------------------
# Markdown / Code / CSV extractors
# ---------------------------------------------------------------------------


def test_markdown_extractor_passthrough(tmp_path: Path) -> None:
    f = tmp_path / "note.md"
    f.write_text("# Title\n\nbody", encoding="utf-8")

    extr = MarkdownExtractor()
    assert extr.can_handle("text/markdown", "md")
    result = extr.extract(f)
    assert "# Title" in result.content_md
    assert result.content_md_truncated is False
    assert result.extractor_used == "markdown"


def test_markdown_extractor_handles_txt(tmp_path: Path) -> None:
    f = tmp_path / "plain.txt"
    f.write_text("plain text content", encoding="utf-8")
    extr = MarkdownExtractor()
    assert extr.can_handle("text/plain", "txt")
    assert "plain text content" in extr.extract(f).content_md


def test_code_extractor_wraps_python_in_fence(tmp_path: Path) -> None:
    f = tmp_path / "snippet.py"
    f.write_text("def hello():\n    return 'hi'\n", encoding="utf-8")

    extr = CodeExtractor()
    assert extr.can_handle(None, "py")
    result = extr.extract(f)
    assert result.content_md.startswith("```python\n")
    assert "def hello():" in result.content_md
    assert result.content_md.rstrip().endswith("```")
    assert result.extractor_used == "code"


def test_code_extractor_maps_cobol(tmp_path: Path) -> None:
    f = tmp_path / "GLMPA013.cbl"
    f.write_text("PROGRAM-ID. GLMPA013.", encoding="utf-8")
    extr = CodeExtractor()
    assert extr.can_handle(None, "cbl")
    assert extr.extract(f).content_md.startswith("```cobol\n")


def test_csv_extractor_5_row_table(tmp_path: Path) -> None:
    f = tmp_path / "small.csv"
    f.write_text(
        "name,age,city\nAlice,30,Paris\nBob,25,Lyon\nCleo,40,Nice\nDan,33,Lille\nEve,28,Caen\n",
        encoding="utf-8",
    )
    extr = CsvExtractor()
    assert extr.can_handle("text/csv", "csv")
    md = extr.extract(f).content_md

    assert "| name | age | city |" in md
    assert "| Alice | 30 | Paris |" in md
    # Separator line.
    assert "|---|---|---|" in md
    # All 5 rows kept (under cap).
    for name in ("Alice", "Bob", "Cleo", "Dan", "Eve"):
        assert name in md


def test_csv_extractor_caps_at_max_rows(tmp_path: Path) -> None:
    rows = ["a,b,c"] + [f"r{i},x,y" for i in range(MAX_ROWS + 50)]
    f = tmp_path / "big.csv"
    f.write_text("\n".join(rows), encoding="utf-8")

    result = CsvExtractor().extract(f)
    assert result.content_md_truncated is True
    assert "more row(s) omitted" in result.content_md
    # The last row inside the cap (index 199 -> "r199") IS present;
    # anything past the cap (e.g. "r205") is omitted.
    assert "| r199 | x | y |" in result.content_md
    assert "| r205 | x | y |" not in result.content_md


def test_csv_extractor_escapes_pipes(tmp_path: Path) -> None:
    f = tmp_path / "tricky.csv"
    f.write_text("a,b\nfoo|bar,baz\n", encoding="utf-8")
    md = CsvExtractor().extract(f).content_md
    # The pipe in "foo|bar" must be escaped so the table renders.
    assert "foo\\|bar" in md


# ---------------------------------------------------------------------------
# PDF extractor (fixture generated in-memory via reportlab)
# ---------------------------------------------------------------------------


def _make_sample_pdf(out: Path, *, pages: int = 2) -> None:
    """Write a tiny multi-page PDF using reportlab."""
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    for i in range(1, pages + 1):
        c.setFont("Helvetica", 14)
        c.drawString(72, 720, f"Page {i} Title")
        c.setFont("Helvetica", 11)
        for j in range(5):
            c.drawString(
                72,
                700 - j * 16,
                f"This is line {j + 1} on page {i} of the sample PDF.",
            )
        c.showPage()
    c.save()
    out.write_bytes(buf.getvalue())


def test_pdf_extractor_basic(tmp_path: Path) -> None:
    f = tmp_path / "sample.pdf"
    _make_sample_pdf(f, pages=2)

    extr = PdfExtractor()
    assert extr.can_handle("application/pdf", "pdf")

    result = extr.extract(f)
    assert result.error is None, f"unexpected error: {result.error}"
    md = result.content_md
    # At least one of our title lines must show up. Both extractors
    # produce slightly different layouts; we check substrings that survive.
    assert "Page 1" in md
    assert "Page 2" in md
    assert result.extractor_used in ("markitdown", "pdfplumber")


# ---------------------------------------------------------------------------
# DOCX extractor (fixture generated via python-docx)
# ---------------------------------------------------------------------------


def _make_sample_docx(out: Path) -> None:
    """Write a tiny DOCX with one heading and two paragraphs."""
    from docx import Document  # type: ignore[import-untyped]

    doc = Document()
    doc.add_heading("Sample Heading", level=1)
    doc.add_paragraph("First paragraph of the sample doc.")
    doc.add_paragraph("Second paragraph with another sentence.")
    doc.save(str(out))


def test_docx_extractor_basic(tmp_path: Path) -> None:
    f = tmp_path / "sample.docx"
    _make_sample_docx(f)

    extr = DocxExtractor()
    assert extr.can_handle(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "docx",
    )

    result = extr.extract(f)
    assert result.error is None, f"unexpected error: {result.error}"
    md = result.content_md
    # MarkItDown emits headings as `# H1`. Be lenient — accept either the
    # markdown form or the plain text containing the heading.
    assert "Sample Heading" in md
    assert "First paragraph" in md
    assert "Second paragraph" in md


# ---------------------------------------------------------------------------
# Registry / select_extractor
# ---------------------------------------------------------------------------


def test_select_extractor_picks_specific_over_generic() -> None:
    """A .csv file goes to CsvExtractor, not CodeExtractor."""
    extr = select_extractor("data.csv", mime_type="text/csv")
    assert isinstance(extr, CsvExtractor)


def test_select_extractor_pdf_by_extension() -> None:
    assert isinstance(select_extractor("report.pdf"), PdfExtractor)


def test_select_extractor_python_by_extension() -> None:
    assert isinstance(select_extractor("script.py"), CodeExtractor)


def test_select_extractor_raises_on_unknown() -> None:
    with pytest.raises(NoExtractorError):
        select_extractor("blob.bin")


def test_extractors_registry_has_no_duplicate_names() -> None:
    """Different extractors may share a backend (markitdown) but each
    class instance must still be unique."""
    types = {type(e) for e in EXTRACTORS}
    assert len(types) == len(EXTRACTORS)
